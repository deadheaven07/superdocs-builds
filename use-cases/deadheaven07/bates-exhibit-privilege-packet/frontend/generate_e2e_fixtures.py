"""Generate realistic sample exhibits for Playwright End-to-End Testing.
Creates mixed format exhibits:
1. Multi-page Contract PDF with SSN, phone, email, names.
2. Confidential Attorney-Client Privilege Memo PDF.
3. Invoice PDF with alphanumeric and numeric account numbers.
4. Scanned Medical Document Image (PNG) with OCR-able text.
5. Case Background Document (DOCX).
"""

import os
from pathlib import Path
import fitz
from PIL import Image, ImageDraw, ImageFont
from docx import Document

FIXTURES_DIR = Path(__file__).parent / "e2e" / "fixtures"
FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

def create_contract_pdf():
    doc = fitz.open()
    
    # Page 1
    page1 = doc.new_page(width=612, height=792)
    lines_p1 = [
        "CONFIDENTIAL MASTER SERVICES AGREEMENT",
        "Exhibit 1 - Contract for Technical Consulting Services",
        "",
        "This Agreement is entered into by and between:",
        "Client: Jane Smith (SSN: 123-45-6789)",
        "Consultant: John Doe (SSN: 987-65-4321)",
        "",
        "Contact Information:",
        "Email: jane.public@example.com",
        "Phone: (212) 555-0199",
        "",
        "Section 1. Scope of Services",
        "The Consultant shall provide digital forensics and litigation data curation.",
        "Both parties agree to hold all proprietary records in strict confidence.",
        "Terms valid from 2026-08-01 to 2027-07-31.",
    ]
    for i, line in enumerate(lines_p1):
        page1.insert_text((72, 80 + i * 22), line, fontsize=11, fontname="helv")
    
    # Page 2
    page2 = doc.new_page(width=612, height=792)
    lines_p2 = [
        "CONFIDENTIAL MASTER SERVICES AGREEMENT (Page 2)",
        "",
        "Section 2. Compensation & Billing Terms",
        "Payment shall be remitted to Account number: ACC-8821-4433",
        "Direct deposit routing account: 8821-4433-2211-9900",
        "",
        "Section 3. Governing Law and Dispute Resolution",
        "This Agreement shall be construed in accordance with the laws of the State of New York.",
        "",
        "Signatures:",
        "Jane Smith, Principal",
        "John Doe, Senior Partner",
    ]
    for i, line in enumerate(lines_p2):
        page2.insert_text((72, 80 + i * 22), line, fontsize=11, fontname="helv")

    path = FIXTURES_DIR / "01_contract_services.pdf"
    doc.save(str(path))
    doc.close()
    print(f"Created: {path}")

def create_privilege_memo_pdf():
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    lines = [
        "CONFIDENTIAL ATTORNEY-CLIENT PRIVILEGED & WORK PRODUCT",
        "MEMORANDUM FOR IN-HOUSE LITIGATION COUNSEL",
        "",
        "DATE: August 14, 2026",
        "TO: Executive Legal Review Committee",
        "FROM: Special Litigation Counsel",
        "RE: Pre-Trial Settlement Valuation and Exposure Analysis",
        "",
        "LEGAL ASSESSMENT:",
        "This privileged legal memorandum summarizes trial strategy, settlement negotiation",
        "thresholds, and risk assessment regarding the pending dispute.",
        "Under Federal Rule of Civil Procedure 26(b)(5), this document constitutes protected",
        "opinion work product and attorney-client communications.",
        "",
        "PROPOSED SETTLEMENT BRACKET: Confidential advisory range $2.4M - $3.1M.",
        "DISCOVERY STATUS: Full privilege withholding requested for exhibit packet.",
    ]
    for i, line in enumerate(lines):
        page.insert_text((72, 80 + i * 22), line, fontsize=11, fontname="helv")
    
    path = FIXTURES_DIR / "02_privileged_strategy_memo.pdf"
    doc.save(str(path))
    doc.close()
    print(f"Created: {path}")

def create_invoice_pdf():
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    lines = [
        "COMMERCIAL STATEMENT & INVOICE",
        "Invoice Number: INV-2026-8841",
        "Invoice Date: 2026-07-20",
        "",
        "Customer Name: Jane Smith",
        "Account Reference: ACC-8821-4433",
        "Secondary Account: ACCOUNT-8821-4433",
        "Wire Transfer Account: 8821-4433-2211-9900",
        "",
        "Itemized Professional Services:",
        "1. Forensic Data Ingestion & Indexing ......... $4,500.00",
        "2. Redaction Review & Validation Audit ......... $2,200.00",
        "3. Cryptographic Manifest Compilation ......... $1,800.00",
        "",
        "TOTAL AMOUNT DUE: $8,500.00",
        "Please remit payment within 30 days.",
    ]
    for i, line in enumerate(lines):
        page.insert_text((72, 80 + i * 22), line, fontsize=11, fontname="helv")
    
    path = FIXTURES_DIR / "03_invoice_billing.pdf"
    doc.save(str(path))
    doc.close()
    print(f"Created: {path}")

def create_scanned_receipt_image():
    # Create high-resolution image to test Tesseract OCR and searchability
    img = Image.new("RGB", (1000, 1400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Draw simple border
    draw.rectangle([(30, 30), (970, 1370)], outline=(180, 180, 180), width=2)
    
    lines = [
        "METROPOLITAN HOSPITAL MEDICAL CLINIC",
        "OFFICIAL CLINICAL TREATMENT RECORD",
        "--------------------------------------------------",
        "Patient Full Name: Jane Smith",
        "Patient ID / SSN: 123-45-6789",
        "Medical Record Account: ACC-9944-1122",
        "Attending Physician: Dr. Robert Vance, MD",
        "Date of Examination: August 05, 2026",
        "",
        "CLINICAL DIAGNOSIS & OBSERVATIONS:",
        "Primary Condition: Acute Bronchitis and Respiratory Infection",
        "Prescription: Amoxicillin 500mg, Oral Suspension",
        "Follow-up: 14 days clinical evaluation",
        "",
        "PATIENT BILLING SUMMARY:",
        "Emergency Room Facility Fee: $1,250.00",
        "Diagnostic Laboratory Panel: $480.00",
        "Pharmacy & Medications: $165.00",
        "--------------------------------------------------",
        "TOTAL CHARGES: $1,895.00",
        "STATUS: PAID VIA INSURANCE CLAIM #MED-9921",
    ]
    
    y = 80
    for line in lines:
        draw.text((70, y), line, fill=(20, 20, 20))
        y += 45
    
    path = FIXTURES_DIR / "04_scanned_medical_receipt.png"
    img.save(str(path))
    print(f"Created: {path}")

def create_docx_memo():
    doc = Document()
    doc.add_heading("Exhibit 5: Technical Background & System Audit Notes", level=1)
    doc.add_paragraph("Author: John Doe, Litigation Discovery Analyst")
    doc.add_paragraph("Subject: E-Discovery Packet Ingestion and Verification Log")
    doc.add_paragraph(
        "This supplementary memorandum documents the chain of custody for exhibits assembled "
        "for court submission. All exhibits have been catalogued and indexed according to their "
        "content rather than arbitrary folder names."
    )
    doc.add_paragraph(
        "Key identifiers in this record include Account number ACC-8821-4433 and consultant "
        "contact jane.public@example.com."
    )
    
    path = FIXTURES_DIR / "05_case_background_notes.docx"
    doc.save(str(path))
    print(f"Created: {path}")

if __name__ == "__main__":
    create_contract_pdf()
    create_privilege_memo_pdf()
    create_invoice_pdf()
    create_scanned_receipt_image()
    create_docx_memo()
    print("All fixtures generated successfully!")
